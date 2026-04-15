import sys
from pathlib import Path
from docx import Document

def extract_questions(docx_path):
    p = Path(docx_path)
    doc = Document(str(p))
    if not doc.tables:
        print(f"[warn] No tables in {p.name}")
        return
    table = doc.tables[0]
    questions = []
    for i, row in enumerate(table.rows):
        text = row.cells[0].text.strip()
        if i == 0 and "question" in text.lower():
            continue
        if text:
            questions.append(text)
    out = p.parent / f"{p.stem}_questions.txt"
    with open(out, "w", encoding="utf-8") as f:
        for n, q in enumerate(questions, 1):
            f.write(f"{n}. {q}\n")
    print(f"[ok] {len(questions)} questions -> {out}")

if __name__ == "__main__":
    extract_questions(sys.argv[1])
