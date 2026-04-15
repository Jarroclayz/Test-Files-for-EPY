# pip install python-docx openai
# ------------------------------------
# Replace YOUR_GITHUB_TOKEN_HERE below with your PAT
GITHUB_TOKEN = "YOUR_GITHUB_TOKEN_HERE"
# ------------------------------------

import os
import re
import threading
from pathlib import Path
import tkinter as tk
from tkinter import filedialog, messagebox
from openai import OpenAI
from docx import Document

client = OpenAI(
    base_url="https://models.inference.ai.azure.com",
    api_key=GITHUB_TOKEN,
)


def get_answers(questions: list[str]) -> list[str]:
    numbered = "\n".join(f"{i+1}. {q}" for i, q in enumerate(questions))
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {
                "role": "system",
                "content": (
                    "You are a helpful assistant. Answer each of the following questions "
                    "concisely and accurately in 1-3 sentences. "
                    "Respond with a numbered list matching the question numbers exactly."
                ),
            },
            {"role": "user", "content": numbered},
        ],
    )
    raw = response.choices[0].message.content.strip()
    # Parse numbered list back into individual answers
    parts = re.split(r"\n?\d+\.\s+", raw)
    answers = [p.strip() for p in parts if p.strip()]
    return answers


def process(filepath: str, status_var: tk.StringVar):
    try:
        status_var.set("Reading questions...")
        doc = Document(filepath)

        if not doc.tables:
            messagebox.showerror("Error", "No table found in the document.")
            status_var.set("Ready.")
            return

        table = doc.tables[0]
        rows = table.rows

        # Detect header row
        start_index = 0
        if "question" in rows[0].cells[0].text.strip().lower():
            start_index = 1

        question_rows = [
            (i, rows[i].cells[0].text.strip())
            for i in range(start_index, len(rows))
            if rows[i].cells[0].text.strip()
        ]

        if not question_rows:
            messagebox.showerror("Error", "No questions found in column 1.")
            status_var.set("Ready.")
            return

        questions = [q for _, q in question_rows]

        status_var.set(f"Getting answers for {len(questions)} questions... please wait ⏳")
        answers = get_answers(questions)

        if len(answers) < len(questions):
            messagebox.showwarning(
                "Warning",
                f"Only {len(answers)} answers returned for {len(questions)} questions. "
                "Some cells may be left blank.",
            )

        status_var.set("Writing answers into document...")
        for idx, (row_index, _) in enumerate(question_rows):
            if idx < len(answers):
                cell = table.rows[row_index].cells[1]
                cell.paragraphs[0].text = answers[idx]

        # Save to Downloads with _completed suffix
        stem = Path(filepath).stem
        downloads = Path.home() / "Downloads"
        out_path = downloads / f"{stem}_completed.docx"
        doc.save(str(out_path))

        status_var.set("Done! ✅")
        messagebox.showinfo("Success! ✅", f"Completed document saved to:\n{out_path}")
        status_var.set("Select a .docx file to get started.")

    except Exception as e:
        messagebox.showerror("Error", f"Something went wrong:\n{e}")
        status_var.set("An error occurred. Please try again.")


def select_file(status_var: tk.StringVar):
    downloads = str(Path.home() / "Downloads")
    filepath = filedialog.askopenfilename(
        title="Select your Word document",
        initialdir=downloads,
        filetypes=[("Word Documents", "*.docx")],
    )
    if not filepath:
        return
    # Run in a background thread so the GUI stays responsive
    thread = threading.Thread(target=process, args=(filepath, status_var), daemon=True)
    thread.start()


# --- Build the GUI ---
root = tk.Tk()
root.title("EPY Auto Answer Generator")
root.geometry("500x250")
root.resizable(False, False)

tk.Label(
    root,
    text="EPY Auto Answer Generator",
    font=("Helvetica", 16, "bold"),
).pack(pady=(30, 5))

tk.Label(
    root,
    text="Select your Word document to get started.",
    font=("Helvetica", 11),
).pack(pady=(0, 20))

status_var = tk.StringVar(value="Select a .docx file to get started.")

tk.Button(
    root,
    text="Select .docx file  ▶",
    font=("Helvetica", 12, "bold"),
    bg="#2C5F8A",
    fg="white",
    padx=10,
    pady=6,
    command=lambda: select_file(status_var),
).pack()

tk.Label(
    root,
    textvariable=status_var,
    font=("Helvetica", 10),
    fg="#555555",
    wraplength=460,
).pack(pady=(20, 0))

root.mainloop()
